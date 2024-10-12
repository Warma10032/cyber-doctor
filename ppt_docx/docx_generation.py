'''将大模型生成的json数据转换为word，可修改代码自定义word的样式'''
import os
import re
import hashlib
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict
from env import get_app_root

_OUTPUT_DIR_DOCX = os.path.join(get_app_root(), "data/cache/docx")

# 如果文件夹路径不存在，先创建
if not os.path.exists(_OUTPUT_DIR_DOCX):
    os.makedirs(_OUTPUT_DIR_DOCX)

def get_file_path_docx(text):
    """生成唯一的文件路径"""
    file_name = hashlib.sha256(text.encode("utf-8")).hexdigest()  # 可以使用uuid替代
    return os.path.join(_OUTPUT_DIR_DOCX, f"{file_name}.docx")

def is_chinese(text: str) -> bool:
    """判断文本中是否包含中文字符"""
    return bool(re.search(r'[\u4e00-\u9fff]', text))

def generate_docx_content(docx_content: Dict) -> str:
    """生成 docx 文件"""
    document = Document()

    # Word 标题
    title_heading = document.add_heading(docx_content['title'], 0)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_heading.runs[0]
    
    # 根据标题是否包含中文设置字体
    if is_chinese(docx_content['title']):
        title_run.font.name = '黑体'  # 中文字体
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 设置中文字体
    else:
        title_run.font.name = 'Arial'  
    
    title_run.font.size = Pt(24)  # 标题字体大小

    # 页内容
    print(f'总共 {len(docx_content["sections"])} 个章节')
    for i, section in enumerate(docx_content['sections']):
        print(f'生成第 {i + 1} 章节: {section["heading"]}')
        section_heading = document.add_heading(section['heading'], level=1)
        section_heading_run = section_heading.runs[0]
        
        # 根据章节标题是否包含中文设置字体
        if is_chinese(section['heading']):
            section_heading_run.font.name = '宋体'
            section_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            section_heading_run.font.name = 'Times New Roman'
        
        section_heading_run.font.size = Pt(16)  # 章节标题字体大小

        for paragraph in section['paragraphs']:
            para_heading = document.add_heading(paragraph['heading'], level=2)
            para_heading_run = para_heading.runs[0]
            
            # 根据段落标题是否包含中文设置字体
            if is_chinese(paragraph['heading']):
                para_heading_run.font.name = '宋体'
                para_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                para_heading_run.font.name = 'Calibri'
            
            para_heading_run.font.size = Pt(14)  # 段落标题字体大小

            # 设置正文内容字体
            p = document.add_paragraph(paragraph['content'])
            p_run = p.runs[0]
            
            # 根据正文内容是否包含中文设置字体
            if is_chinese(paragraph['content']):
                p_run.font.name = '宋体'
                p_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                p_run.font.name = 'Arial'
            
            p_run.font.size = Pt(12)  # 正文字体大小

    _output_file = get_file_path_docx(str(time.time()))
    document.save(_output_file)

    return _output_file

